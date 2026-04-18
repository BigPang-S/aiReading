#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


SPACE_RE = re.compile(r"\s+")


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="把纯文本章节渲染成规范化 docx")
    parser.add_argument("--input", required=True, help="输入 txt 文件")
    parser.add_argument("--output", required=True, help="输出 docx 文件")
    return parser.parse_args()


def set_run_font(run, font_name: str, size_pt: float, *, bold: bool = False) -> None:
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for key in ("w:eastAsia", "w:ascii", "w:hAnsi"):
        rfonts.set(qn(key), font_name)


def build_docx(title: str, paragraphs: list[str], output_path: Path) -> None:
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    normal = doc.styles["Normal"]
    normal.font.name = "宋体"
    normal.font.size = Pt(12)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    normal.paragraph_format.line_spacing = Pt(29)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.paragraph_format.space_before = Pt(0)
    title_para.paragraph_format.space_after = Pt(0)
    title_run = title_para.add_run(title)
    set_run_font(title_run, "黑体", 16)

    for text in paragraphs:
        para = doc.add_paragraph()
        para.paragraph_format.first_line_indent = Pt(24)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        para.paragraph_format.line_spacing = Pt(29)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0)
        run = para.add_run(text)
        set_run_font(run, "宋体", 12)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


def load_text_chapter(input_path: Path) -> tuple[str, list[str]]:
    lines = [line.rstrip() for line in input_path.read_text(encoding="utf-8").splitlines()]
    non_empty = [line for line in lines if line.strip()]
    if not non_empty:
        raise ValueError(f"输入文件为空：{input_path}")

    title = non_empty[0].strip()
    body_lines = lines[lines.index(non_empty[0]) + 1 :]

    paragraphs: list[str] = []
    bucket: list[str] = []
    for line in body_lines:
        stripped = line.strip()
        if not stripped:
            if bucket:
                paragraphs.append("".join(bucket))
                bucket = []
            continue
        bucket.append(stripped)
    if bucket:
        paragraphs.append("".join(bucket))
    return title, paragraphs


def count_body_chars(paragraphs: list[str]) -> int:
    return len(SPACE_RE.sub("", "".join(paragraphs)))


def main() -> int:
    args = parse_args()
    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    title, paragraphs = load_text_chapter(input_path)
    build_docx(title, paragraphs, output_path)
    print(f"已生成：{output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
